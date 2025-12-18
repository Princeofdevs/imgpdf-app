import os
import uuid
import math
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
from werkzeug.utils import secure_filename
from PIL import Image
import img2pdf
from pdf2image import convert_from_path

# ======================
# CONFIGURATION
# ======================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
CONVERTED_FOLDER = os.path.join(BASE_DIR, "converted")

ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg"}
ALLOWED_PDF_EXTENSIONS = {"pdf"}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

app = Flask(__name__, template_folder="templates", static_folder="static")
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["CONVERTED_FOLDER"] = CONVERTED_FOLDER
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_SIZE
app.secret_key = "dev-secret-key"

# ======================
# HELPERS
# ======================
def allowed_file(filename, allowed):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed

def unique_name(filename):
    return f"{uuid.uuid4().hex}_{secure_filename(filename)}"

# ======================
# CHECK IF python-docx IS INSTALLED
# ======================
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ModuleNotFoundError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Image → Word feature will be disabled.")


# ======================
# ROUTES
# ======================
@app.route("/")
def index():
    return render_template("index.html", docx_available=DOCX_AVAILABLE)


# ----------------------
# IMAGE ➜ PDF
# ----------------------
@app.route("/image-to-pdf", methods=["POST"])
def image_to_pdf():
    if "images" not in request.files:
        flash("No images uploaded")
        return redirect(url_for("index"))

    files = request.files.getlist("images")
    saved_images = []

    for file in files:
        if file.filename == "":
            continue
        if not allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
            flash("Only PNG, JPG, JPEG allowed")
            return redirect(url_for("index"))

        path = os.path.join(UPLOAD_FOLDER, unique_name(file.filename))
        file.save(path)
        saved_images.append(path)

    if not saved_images:
        flash("No valid images")
        return redirect(url_for("index"))

    layout = request.form.get("layout", "single")
    temp_files = []
    final_images = []

    IMAGES_PER_PAGE = 4 if layout == "multiple" else 1
    chunks = [saved_images[i:i + IMAGES_PER_PAGE] for i in range(0, len(saved_images), IMAGES_PER_PAGE)]

    try:
        for chunk in chunks:
            if len(chunk) == 1:
                final_images.append(chunk[0])
                continue

            cols = 2
            rows = math.ceil(len(chunk) / cols)
            imgs = [Image.open(p).convert("RGB") for p in chunk]

            max_w = max(i.width for i in imgs)
            max_h = max(i.height for i in imgs)

            canvas = Image.new("RGB", (cols * max_w, rows * max_h), (255, 255, 255))
            for i, img in enumerate(imgs):
                r = i // cols
                c = i % cols
                x = c * max_w + (max_w - img.width) // 2
                y = r * max_h + (max_h - img.height) // 2
                canvas.paste(img, (x, y))

            stitched = os.path.join(UPLOAD_FOLDER, f"stitched_{uuid.uuid4().hex}.jpg")
            canvas.save(stitched)
            final_images.append(stitched)
            temp_files.append(stitched)

    except Exception as e:
        flash(f"Image processing failed: {e}")
        return redirect(url_for("index"))

    pdf_name = f"converted_{uuid.uuid4().hex}.pdf"
    pdf_path = os.path.join(CONVERTED_FOLDER, pdf_name)

    try:
        with open(pdf_path, "wb") as f:
            f.write(img2pdf.convert(final_images))
    except Exception as e:
        flash(f"PDF generation failed: {e}")
        return redirect(url_for("index"))
    finally:
        for f in saved_images + temp_files:
            if os.path.exists(f):
                os.remove(f)

    return render_template("index.html", pdf_download_link=pdf_name, docx_available=DOCX_AVAILABLE)


# ----------------------
# IMAGE ➜ WORD (only if docx available)
# ----------------------
if DOCX_AVAILABLE:
    @app.route("/image-to-word", methods=["POST"])
    def image_to_word():
        if "images_word" not in request.files:
            flash("No images uploaded")
            return redirect(url_for("index"))

        files = request.files.getlist("images_word")
        saved_images = []

        for file in files:
            if file.filename == "":
                continue
            if not allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
                flash("Only PNG, JPG, JPEG allowed")
                return redirect(url_for("index"))
            path = os.path.join(UPLOAD_FOLDER, unique_name(file.filename))
            file.save(path)
            saved_images.append(path)

        if not saved_images:
            flash("No valid images")
            return redirect(url_for("index"))

        doc = Document()
        for img_path in saved_images:
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph("")

        word_name = f"converted_{uuid.uuid4().hex}.docx"
        word_path = os.path.join(CONVERTED_FOLDER, word_name)
        doc.save(word_path)

        for f in saved_images:
            if os.path.exists(f):
                os.remove(f)

        return render_template("index.html", word_download_link=word_name, docx_available=DOCX_AVAILABLE)


# ----------------------
# PDF ➜ IMAGE
# ----------------------
@app.route("/pdf-to-image", methods=["POST"])
def pdf_to_image():
    if "pdf" not in request.files:
        flash("No PDF uploaded")
        return redirect(url_for("index"))

    file = request.files["pdf"]

    if not allowed_file(file.filename, ALLOWED_PDF_EXTENSIONS):
        flash("Only PDF allowed")
        return redirect(url_for("index"))

    pdf_path = os.path.join(UPLOAD_FOLDER, unique_name(file.filename))
    file.save(pdf_path)

    try:
        images = convert_from_path(pdf_path, output_folder=CONVERTED_FOLDER, fmt="jpeg")
        image_files = [os.path.basename(i.filename) for i in images]
    except Exception as e:
        flash(f"PDF conversion failed: {e}")
        return redirect(url_for("index"))
    finally:
        os.remove(pdf_path)

    return render_template("index.html", image_download_links=image_files, docx_available=DOCX_AVAILABLE)


# ----------------------
# DOWNLOAD
# ----------------------
@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory(
        CONVERTED_FOLDER,
        filename,
        as_attachment=True
    )


# ======================
# RUN APP
# ======================
if __name__ == "__main__":
    app.run()
